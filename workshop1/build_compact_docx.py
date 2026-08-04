"""Build compact3 docx with figures from compact3_draft.md."""
import re, base64
from pathlib import Path
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt

REPO = Path(r"C:/Users/soham/AppData/Local/hermes/research-corpus/precision-cliff")
OUT = Path(__file__).parent
FIGS = OUT / "figs"
FIGS.mkdir(exist_ok=True)

# ---------- Figure 1: predicted vs best-in-family ----------
N = [13, 17, 21, 31, 35, 37, 43]
pred = [1.6250000, 2.0517767, 2.1000000, 2.5833333, 2.9166667, 3.0345178, 3.0714286]
best = [1.7761424, 2.0517767, 2.2588835, 2.7485281, 2.9166667, 3.0345178, 3.2416246]
modal = [10/18, 3/4, 12/15, 13/17, 3/4, 3/4, 6/7]

fig, ax = plt.subplots(figsize=(5.2, 3.0), dpi=72)
ax.plot(N, best, "o--", color="#888888", label="best in recipe family", lw=1.4, ms=5)
ax.plot(N, pred, "s-", color="#1f5fa8", label="predicted = empirical mode (7/7)", lw=1.6, ms=6)
for n, p, b in zip(N, pred, best):
    if b - p > 1e-9:
        ax.annotate(f"+{b-p:.3f}", xy=(n, (p+b)/2), fontsize=7.5, color="#a03030",
                    ha="left", xytext=(n+0.4, (p+b)/2 - 0.01))
        ax.vlines(n, p, b, color="#a03030", lw=1, alpha=0.6)
ax.set_xlabel("N (circles)")
ax.set_ylabel("sum of radii")
ax.set_title("Closed-form prediction vs best-in-family, 7 held-out cells", fontsize=10)
ax.legend(fontsize=8, loc="lower right")
ax.grid(alpha=0.25, lw=0.5)
fig.tight_layout()
fig.savefig(FIGS / "fig1.png")
plt.close(fig)

# ---------- Figure 2: tier ladder + on-prediction ----------
tiers = ["weak\n(haiku)", "middle\n(sonnet)", "top\n(opus_alias\u2021)"]
validity = [83, 100, 13]      # matched cells, 10^-6
onpred = [70, 3.3, 0]         # 35/50, 1/30, 0/4
x = range(3)
fig, ax = plt.subplots(figsize=(4.6, 2.8), dpi=72)
w = 0.38
b1 = ax.bar([i - w/2 for i in x], validity, w, color="#1f5fa8", label="valid at 10\u207b\u2076 (%)")
b2 = ax.bar([i + w/2 for i in x], onpred, w, color="#d08a2e", label="on-prediction (%)")
for r in list(b1) + list(b2):
    ax.annotate(f"{r.get_height():.0f}", (r.get_x() + r.get_width()/2, r.get_height()),
                ha="center", va="bottom", fontsize=8)
ax.set_xticks(list(x)); ax.set_xticklabels(tiers, fontsize=9)
ax.set_ylim(0, 112)
ax.set_ylabel("percent of invocations")
ax.set_title("Tier ladder, matched cells N \u2208 {13, 21, 31}", fontsize=10)
ax.legend(fontsize=8)
ax.grid(axis="y", alpha=0.25, lw=0.5)
fig.tight_layout()
fig.savefig(FIGS / "fig2.png")
plt.close(fig)

# ---------- docx from markdown ----------
md = (REPO / "workshop1" / "compact3_draft.md").read_text(encoding="utf-8")
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(10)

def add_runs(par, text):
    # split on **bold** ; strip markdown escapes \* and ` and *
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        bold = part.startswith("**") and part.endswith("**")
        t = part[2:-2] if bold else part
        t = t.replace("\\*", "*").replace("`", "")
        # italics: single *...*
        subs = re.split(r"(?<!\*)\*(?!\*)([^*]+)\*(?!\*)", t)
        for i, s in enumerate(subs):
            r = par.add_run(s)
            r.bold = bold
            if i % 2 == 1:
                r.italic = True

lines = md.splitlines()
i = 0
para_buf = []
def flush():
    global para_buf
    if para_buf:
        p = doc.add_paragraph()
        add_runs(p, " ".join(para_buf))
        para_buf = []

fig_inserted = {1: False, 2: False}
while i < len(lines):
    ln = lines[i]
    if ln.startswith("| N | predicted"):
        flush()
        # parse table block
        rows = []
        while i < len(lines) and lines[i].startswith("|"):
            cells = [c.strip() for c in lines[i].strip("|").split("|")]
            rows.append(cells)
            i += 1
        rows = [r for r in rows if not set("".join(r)) <= set("-: ")]
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for ri, r in enumerate(rows):
            for ci, c in enumerate(r):
                t.cell(ri, ci).text = c.replace("\\*", "*")
                for p_ in t.cell(ri, ci).paragraphs:
                    for run in p_.runs:
                        run.font.size = Pt(9)
                        if ri == 0:
                            run.bold = True
        continue
    if ln.startswith("## "):
        flush()
        # insert fig1 right after section 2 heading paragraph content ends -> handled below
        doc.add_heading(ln[3:].replace("\\*", "*"), level=1)
        i += 1
        continue
    if ln.startswith("# "):
        flush()
        doc.add_heading(ln[2:].replace("\\*", "*"), level=0)
        i += 1
        continue
    if ln.strip() == "---":
        flush()
        i += 1
        continue
    if ln.strip() == "":
        flush()
        # figure hooks: after the table paragraph block in section 2, and after tier ladder section
        i += 1
        continue
    para_buf.append(ln.strip())
    i += 1
flush()

# Re-walk: insert figures after known anchor paragraphs
from docx.text.paragraph import Paragraph
import copy

def insert_fig_after(anchor_snippet, png, caption, width=5.6):
    for idx, par in enumerate(doc.paragraphs):
        if anchor_snippet in par.text:
            # add image paragraph after this one
            new_p = par._p.addnext(copy.deepcopy(doc.paragraphs[0]._p))
            p = Paragraph(par._p.getnext(), par._parent)
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            run = p.add_run()
            run.add_picture(str(png), width=Inches(width))
            cap_el = p._p.addnext(copy.deepcopy(doc.paragraphs[0]._p))
            cp = Paragraph(p._p.getnext(), p._parent)
            for r in list(cp.runs):
                r._r.getparent().remove(r._r)
            cr = cp.add_run(caption)
            cr.italic = True
            cr.font.size = Pt(8.5)
            return True
    return False

ok1 = insert_fig_after("the closed form is a weak-tier regularity", FIGS / "fig1.png",
    "Figure 1: Predicted value (= empirical modal output at all 7 N) against the best value available "
    "inside the recipe family. Red segments mark the four cells where the emitted template leaves value "
    "on the table (gaps +0.151 to +0.170); at N = 17, 35, 37 the prediction is family-optimal.")
ok2 = insert_fig_after("The ladder:", FIGS / "fig2.png",
    "Figure 2: Ambition rises with tier while validity does not. Matched cells; on-prediction: 35/50, 1/30, 0/4. "
    "\u2021 opus_alias = unattributable serving alias, caveat as in text.", width=4.8)

out = OUT / "compact3_figs.docx"
doc.save(out)
print("fig anchors:", ok1, ok2)
print("docx bytes:", out.stat().st_size)
b64 = base64.b64encode(out.read_bytes()).decode()
(OUT / "compact3_figs.docx.b64").write_text(b64)
print("b64 chars:", len(b64))
